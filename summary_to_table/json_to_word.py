from docx import Document
# from docx2pdf import convert
import json
from docx import Document as DocxDocument
# from docx2pdf import convert
import json
from spire.doc import *
from spire.doc.common import *






def load_summaries(json_path):
    try:
        with open(json_path, 'r') as f:
            return json.load(f)
    except (json.JSONDecodeError, FileNotFoundError):
        return []

# Usage:

def json_to_word():
    raw = load_summaries('cache/summaries.json')
    if isinstance(raw, dict):
        summaries = [v["summary"] if "summary" in v else v for v in raw.values()]
    else:
        summaries = raw
    output_path = "paper_summary.docx"
    doc = DocxDocument()
    doc.add_heading("All Research Paper Summaries", level=1)
    if len(summaries) > 0:
    
        for idx, summary in enumerate(summaries, start=1):
            doc.add_heading(f"Summary {idx}", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Field"
            hdr_cells[1].text = "Details"
            for key, value in summary.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                if isinstance(value, list) and value and isinstance(value[0], dict):
                    nested_table = row_cells[1].add_table(rows=1, cols=len(value[0]))
                    hdr_cells = nested_table.rows[0].cells
                    for i, k in enumerate(value[0].keys()):
                        hdr_cells[i].text = str(k)
                    for item in value:
                        row = nested_table.add_row().cells
                        for i, k in enumerate(item.keys()):
                            row[i].text = str(item[k])
                elif isinstance(value, list):
                    row_cells[1].text = ", ".join(str(v) for v in value)
                elif isinstance(value, dict):
                    nested_table = row_cells[1].add_table(rows=1, cols=2)
                    hdr_cells = nested_table.rows[0].cells
                    hdr_cells[0].text = "Key"
                    hdr_cells[1].text = "Value"
                    for k, v in value.items():
                        row = nested_table.add_row().cells
                        row[0].text = str(k)
                        row[1].text = str(v)
                else:
                    row_cells[1].text = str(value)
            doc.add_paragraph("")
        doc.save(output_path)
        # convert("paper_summary.docx", "paper_summary.pdf")
        

        document = Document()
        document.LoadFromFile("paper_summary.docx")
        document.SaveToFile("paper_summary.pdf", FileFormat.PDF)
        document.Close()

# json_to_word_table(paper["summary"], "paper_summary.docx")
